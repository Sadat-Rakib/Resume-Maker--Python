from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# The following code creates a new Word document
doc = Document()

# The following function code is used to add headings
def add_heading(text, size=16):
    para = doc.add_paragraph()
    run = para.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = "Rubik"
    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# The following function code is used to add body text.
def add_body_text(text, size=12, bold=False, italic=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Rubik"
    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# The following code adds the header: Name and Job Title
p = doc.add_paragraph()
name_run = p.add_run("MIR SADAT BIN RAKIB\n")
name_run.bold = True
name_run.font.size = Pt(24)
name_run.font.name = "Rubik"

title_run = p.add_run("Front End Developer")
title_run.bold = True
title_run.font.size = Pt(14)
p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# The following code adds the contact info
add_body_text("📞 +12638810363 | ✉ mirsadatbinrakib01@gmail.com | 🌐 linkedin.com/in/mir-sadat-bin-rakib\n", size=11)

# The following code adds the Projects Section
add_heading("Projects")
add_body_text("Fantastic.Furniture.Shop | Developed a modern, responsive furniture website using HTML, CSS & JS, enhancing product visibility by 40% and improving customer interaction by 30%.\n", bold=True)
add_body_text("Special Restaurant | Built a fully responsive restaurant website using HTML, CSS, JavaScript & Bootstrap, improving reservations by 35% and increasing user engagement by 25%.\n", bold=True)

# The following code adds the Experience Section
add_heading("Experience")
add_body_text("Teacher's Assistant | Scholastica\n07/2023 - 04/2024 | Dhaka, Bangladesh\n", bold=True, italic=True)
add_body_text("• Developed web-based projects for classroom activities, increasing student engagement by 40%.\n• Mentored and assisted 50+ students in understanding programming concepts.\n")

# The following code adds the Skills Section
add_heading("Skills")
add_body_text("HTML, CSS, JavaScript, Bootstrap, React.js, Typescript\nNext.js, Figma, Git, GitHub, Vercel, Python, Responsive Web\nProblem-Solving, Adaptability, Continuous Learning\n", bold=True)

# Finally you use the followijng code to save your precious resume document
doc.save("Resume_Template.docx")
print("Resume_Template.docx has been created successfully!")
